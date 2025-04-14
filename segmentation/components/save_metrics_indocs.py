# save_metrics_indocs.py

import streamlit as st
import os
import json
from docx import Document
from segmentation.constant.config import TRAIN_METRIC_DIRR, TEST_METRIC_DIRR
from segmentation.entity.config_entity import ArtifactConfig

# Instantiate artifact config to get the model base path
artifact_config = ArtifactConfig()
model_base_path = artifact_config.trained_models_dirr

def format_val(val):
    """Format numeric values to 4 decimal places, else str."""
    if isinstance(val, (float, int)):
        return f"{val:.4f}"
    return str(val)

def read_json_file(filepath):
    """Read JSON, return dict/list or None on error."""
    try:
        with open(filepath, "r") as f:
            return json.load(f)
    except Exception as e:
        st.error(f"Error reading {filepath}: {e}")
        return None

def process_train_metrics(train_data):
    """
    From a list of epoch records, pick the one with highest valid_metrics['iou_score'].
    Also sum total_time across all epochs.
    Returns: best_epoch, best_record, total_time
    """
    best_epoch = None
    best_val_iou = -1.0
    best_record = None
    total_time = 0.0

    for record in train_data:
        epoch = record.get("epoch")
        valid = record.get("valid_metrics", {})
        val_iou = valid.get("iou_score", 0.0)
        if val_iou > best_val_iou:
            best_val_iou = val_iou
            best_epoch = epoch
            best_record = record

        train = record.get("train_metrics", {})
        total_time += train.get("total_time", 0.0)

    return best_epoch, best_record, total_time

def create_doc_table_all(model_base_path):
    """Loop over datasets/models, write a .docx report with best-epoch metrics."""
    document = Document()
    document.add_heading("Model Metrics Report", level=1)

    # find all dataset dirs
    dataset_dirs = [
        d for d in os.listdir(model_base_path)
        if os.path.isdir(os.path.join(model_base_path, d))
    ]
    if not dataset_dirs:
        document.add_paragraph("No datasets found in the model base path.")
        out = "metrics_report_all.docx"
        document.save(out)
        return out

    for dataset in dataset_dirs:
        document.add_heading(f"Dataset: {dataset}", level=2)
        dataset_path = os.path.join(model_base_path, dataset)
        model_dirs = [
            m for m in os.listdir(dataset_path)
            if os.path.isdir(os.path.join(dataset_path, m))
        ]
        if not model_dirs:
            document.add_paragraph("No models found for this dataset.")
            continue

        for model in model_dirs:
            document.add_heading(f"Model: {model}", level=3)

            # paths to JSON metric files
            train_path = os.path.join(dataset_path, model, TRAIN_METRIC_DIRR)
            test_path  = os.path.join(dataset_path, model, TEST_METRIC_DIRR)

            train_data = read_json_file(train_path)
            test_data  = read_json_file(test_path)

            # --- Train & Validation (Best Epoch) ---
            if train_data:
                best_epoch, best_record, total_time = process_train_metrics(train_data)
                document.add_paragraph(
                    "Training & Validation Metrics for Best Epoch (by validation IoU):"
                )
                table = document.add_table(rows=1, cols=13)
                hdr = table.rows[0].cells
                hdr[0].text  = "Epoch"
                hdr[1].text  = "Train Dice Loss"
                hdr[2].text  = "Train IoU"
                hdr[3].text  = "Train Accuracy"
                hdr[4].text  = "Train Precision"
                hdr[5].text  = "Train Recall"
                hdr[6].text  = "Train Fscore"
                hdr[7].text  = "Valid Dice Loss"
                hdr[8].text  = "Valid IoU"
                hdr[9].text  = "Valid Accuracy"
                hdr[10].text = "Valid Precision"
                hdr[11].text = "Valid Recall"
                hdr[12].text = "Valid Fscore"

                # fill row
                train_m = best_record.get("train_metrics", {})
                valid_m = best_record.get("valid_metrics", {})
                row = table.add_row().cells
                row[0].text  = format_val(best_epoch)
                row[1].text  = format_val(train_m.get("dice_loss", "N/A"))
                row[2].text  = format_val(train_m.get("iou_score", "N/A"))
                row[3].text  = format_val(train_m.get("accuracy", "N/A"))
                row[4].text  = format_val(train_m.get("precision", "N/A"))
                row[5].text  = format_val(train_m.get("recall", "N/A"))
                row[6].text  = format_val(train_m.get("fscore", "N/A"))
                row[7].text  = format_val(valid_m.get("dice_loss", "N/A"))
                row[8].text  = format_val(valid_m.get("iou_score", "N/A"))
                row[9].text  = format_val(valid_m.get("accuracy", "N/A"))
                row[10].text = format_val(valid_m.get("precision", "N/A"))
                row[11].text = format_val(valid_m.get("recall", "N/A"))
                row[12].text = format_val(valid_m.get("fscore", "N/A"))

                document.add_paragraph(f"Best Epoch: {best_epoch}")
                document.add_paragraph(f"Total Training Time: {format_val(total_time)} seconds")
            else:
                document.add_paragraph("No training metrics found.")

            # --- Test Metrics ---
            if test_data:
                document.add_paragraph("Test Metrics:")
                ttable = document.add_table(rows=1, cols=6)
                th = ttable.rows[0].cells
                th[0].text = "Dice Loss"
                th[1].text = "IoU Score"
                th[2].text = "Accuracy"
                th[3].text = "Precision"
                th[4].text = "Recall"
                th[5].text = "Fscore"

                if isinstance(test_data, list) and test_data:
                    tm = test_data[0].get("test_metrics", {})
                    r = ttable.add_row().cells
                    r[0].text = format_val(tm.get("dice_loss", "N/A"))
                    r[1].text = format_val(tm.get("iou_score", "N/A"))
                    r[2].text = format_val(tm.get("accuracy", "N/A"))
                    r[3].text = format_val(tm.get("precision", "N/A"))
                    r[4].text = format_val(tm.get("recall", "N/A"))
                    r[5].text = format_val(tm.get("fscore", "N/A"))
                else:
                    document.add_paragraph("Test metrics data is empty.")
            else:
                document.add_paragraph("No test metrics found.")

            document.add_paragraph("-----------------------------")
        document.add_paragraph("====================================")

    output_path = "metrics_report_all.docx"
    document.save(output_path)
    return output_path

def create_doc_table_tab():
    """Streamlit UI: button to trigger the report creation."""
    st.header("Create Doc Table")
    st.write("Click the Start button to generate the metrics report.")

    if st.button("Start"):
        output_docx = create_doc_table_all(model_base_path)
        st.success(f"Metrics report created: {output_docx}")
        with open(output_docx, "rb") as file:
            st.download_button(
                "Download Report",
                data=file,
                file_name=output_docx,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
