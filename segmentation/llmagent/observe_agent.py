import os
import json
import re
from docx import Document
from agno.agent import Agent
from agno.models.ollama import Ollama
from agno.tools.pandas import PandasTools

from segmentation.entity.config_entity import ArtifactConfig
from segmentation.constant.config import TRAIN_METRIC_DIRR, TEST_METRIC_DIRR

# Initialize the review Agent with PandasTools
review_agent = Agent(
    model=Ollama(id="llama3.2"),
    tools=[PandasTools()],
    show_tool_calls=True
)

# Get the base model directory from configuration
artifact_config = ArtifactConfig()
model_path = artifact_config.trained_models_dirr

# List all model directories
models = [m for m in os.listdir(model_path) if os.path.isdir(os.path.join(model_path, m))]

def load_metrics(json_path):
    """Load JSON data from the given file path."""
    if os.path.isfile(json_path):
        with open(json_path, 'r') as f:
            data = json.load(f)
        return data
    return []

def compute_model_summary(metrics_data):
    """
    Given a list of epoch data, compute a summary including:
      - number of epochs
      - final epoch's train and validation metrics (all metrics)
      - best epoch (based on highest validation accuracy) and its metrics
      - total training time (summing the total_time from train_metrics in each epoch)
      - average train and validation accuracy (for reference)
    """
    if not metrics_data:
        return {}
    
    num_epochs = len(metrics_data)
    final_epoch = metrics_data[-1]
    final_train = final_epoch.get("train_metrics", {})
    final_valid = final_epoch.get("valid_metrics", {})
    
    # Initialize best as final epoch and accumulate total_time and accuracies.
    best_epoch_num = final_epoch.get("epoch", num_epochs)
    best_valid_acc = final_valid.get("accuracy", 0)
    best_train = final_train  # best train metrics corresponding to best valid
    total_time = 0
    total_train_acc = 0
    total_valid_acc = 0
    
    for epoch in metrics_data:
        train = epoch.get("train_metrics", {})
        valid = epoch.get("valid_metrics", {})
        total_time += train.get("total_time", 0)  # total_time is inside train_metrics
        total_train_acc += train.get("accuracy", 0)
        total_valid_acc += valid.get("accuracy", 0)
        if valid.get("accuracy", 0) > best_valid_acc:
            best_valid_acc = valid.get("accuracy", 0)
            best_epoch_num = epoch.get("epoch", None)
            best_train = train
            final_valid = valid  # update final_valid as best valid metrics
    
    avg_train_acc = total_train_acc / num_epochs
    avg_valid_acc = total_valid_acc / num_epochs

    summary = {
        "num_epochs": num_epochs,
        "final_train": final_train,
        "final_valid": final_valid,
        "best_epoch": best_epoch_num,
        "avg_train_acc": avg_train_acc,
        "avg_valid_acc": avg_valid_acc,
        "total_time": total_time
    }
    return summary

def format_metrics(metrics):
    """
    Given a dictionary of metrics, return a formatted string
    where each metric is on a new line with format:
      key: value
    """
    lines = []
    for key, value in metrics.items():
        # Skip total_time here if you want to show it separately
        if key == "total_time":
            continue
        lines.append(f"{key}: {value:.4f}" if isinstance(value, (int, float)) else f"{key}: {value}")
    return "\n".join(lines)

def add_run_markdown(paragraph, text):
    """
    Process inline markdown for bold text.
    Splits text by ** and makes those runs bold.
    """
    segments = re.split(r'(\*\*.*?\*\*)', text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
        else:
            paragraph.add_run(seg)

def add_markdown_to_doc(doc, markdown_text):
    """
    A simple markdown parser that adds headings, bullet lists, and bold text to the docx Document.
    - Lines starting with "### " become level-3 headings.
    - Lines starting with "- " or "* " become bullet list items.
    - Other lines are added as normal paragraphs with inline bold formatting.
    """
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("* ") or stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_markdown(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_run_markdown(p, stripped)

def create_consolidated_report(model_summaries):
    """
    Create one consolidated DOCX report that includes:
      - For each model: a dedicated table with its metrics overview.
      - A training time table across all models.
      - An overall performance review (aggregated from all models).
    """
    doc = Document()
    doc.add_heading("Consolidated Model Performance Report", 0)

    # === Per-Model Metrics Overview ===
    doc.add_heading("Per-Model Metrics Overview", level=1)
    for summary in model_summaries:
        doc.add_heading(summary["model_name"], level=2)
        # Create a table with 2 columns: Metric and Value.
        table = doc.add_table(rows=0, cols=2)
        # Add rows for:
        #   - Number of Epochs
        #   - Final Train Metrics (multiline cell)
        #   - Final Valid Metrics (multiline cell)
        #   - Best Epoch (with best valid accuracy)
        #   - Total Training Time (sum)
        #   - Average Train and Valid Accuracy
        def add_row(label, value):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
        
        add_row("# Epochs", summary["num_epochs"])
        add_row("Final Train Metrics", format_metrics(summary["final_train"]))
        add_row("Final Valid Metrics", format_metrics(summary["final_valid"]))
        add_row("Best Epoch", summary["best_epoch"])
        add_row("Total Training Time (sec)", f'{summary["total_time"]:.2f}')
        add_row("Avg Train Accuracy", f'{summary["avg_train_acc"]:.4f}')
        add_row("Avg Valid Accuracy", f'{summary["avg_valid_acc"]:.4f}')
        doc.add_paragraph()  # blank line after each model table

    # === Overall Training Time Table ===
    doc.add_heading("Training Time Overview", level=1)
    time_table = doc.add_table(rows=1, cols=2)
    hdr_cells = time_table.rows[0].cells
    hdr_cells[0].text = "Model"
    hdr_cells[1].text = "Total Training Time (sec)"
    for summary in model_summaries:
        row_cells = time_table.add_row().cells
        row_cells[0].text = summary["model_name"]
        row_cells[1].text = f'{summary["total_time"]:.2f}'
    doc.add_paragraph()

    # === Overall Performance Review ===
    doc.add_heading("Overall Performance Review", level=1)
    # Identify best model overall based on highest final validation accuracy.
    best_model = max(model_summaries, key=lambda s: s["final_valid"].get("accuracy", 0))
    overall_data = {
        "models": [
            {
                "model_name": s["model_name"],
                "final_valid": s["final_valid"],
                "best_epoch": s["best_epoch"],
                "total_time": s["total_time"]
            } for s in model_summaries
        ],
        "best_model": best_model["model_name"],
        "review_notes": (
            "Provide an overall performance analysis based on the above metrics. "
            "Include insights on the final performance metrics, best epochs, total training times, "
            "and recommendations on which model performed best and why."
        )
    }
    review_query = f"Review the following aggregated model performance data and provide a detailed overall analysis:\n{overall_data}"
    response = review_agent.run(review_query)
    overall_review = response.content
    add_markdown_to_doc(doc, overall_review)
    
    # Save the consolidated report in the base model directory.
    report_path = os.path.join(model_path, "Consolidated_Model_Performance_Report.docx")
    doc.save(report_path)
    print(f"✅ Consolidated Report created: {report_path}")

def start_agent():
    model_summaries = []
    
    # Process each model directory.
    for model_name in models:
        print(f"\n📂 Processing Model: {model_name}")
        # Assume training JSON file holds multi-epoch data.
        train_metrics_path = os.path.join(model_path, model_name, TRAIN_METRIC_DIRR)
        metrics_data = load_metrics(train_metrics_path)
        summary = compute_model_summary(metrics_data)
        summary["model_name"] = model_name
        model_summaries.append(summary)
    
    # Create the consolidated DOCX report.
    create_consolidated_report(model_summaries)
